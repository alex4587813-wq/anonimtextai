from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_DIR = Path(__file__).resolve().parent.parent
OUTPUT_DIR = PROJECT_DIR / "output" / "documents"
OUTPUT_PATH = OUTPUT_DIR / "Инструкция пользователя — Локальный анонимизатор.docx"

FONT = "Calibri"
BODY_COLOR = RGBColor(34, 34, 34)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(100, 100, 100)
LIGHT_FILL = "F4F6F9"
TABLE_HEADER_FILL = "E8EEF5"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, size=11, color=BODY_COLOR, bold=None, italic=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in tokens.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_numbering_definition(doc, num_format, text, font=None):
    numbering = doc.part.numbering_part.element
    existing_abstract = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(existing_abstract, default=-1) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), num_format)
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), text)
    level.append(level_text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)

    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    ppr.append(indent)
    level.append(ppr)

    if font:
        rpr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), font)
        fonts.set(qn("w:hAnsi"), font)
        rpr.append(fonts)
        level.append(rpr)

    abstract.append(level)
    numbering.append(abstract)

    existing_num = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    num_id = max(existing_num, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25


def add_bullet(doc, text, bullet_num_id):
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, bullet_num_id)
    set_run_font(paragraph.add_run(text))
    return paragraph


def add_step(doc, title, detail, decimal_num_id):
    paragraph = doc.add_paragraph()
    apply_numbering(paragraph, decimal_num_id)
    title_run = paragraph.add_run(f"{title}. ")
    set_run_font(title_run, bold=True)
    set_run_font(paragraph.add_run(detail))
    return paragraph


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_cell_text(cell, text, bold=False, color=BODY_COLOR, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=color, bold=bold)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)

    header_row_properties = table.rows[0]._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    header_row_properties.append(table_header)

    for index, header in enumerate(headers):
        shade_cell(table.rows[0].cells[index], TABLE_HEADER_FILL)
        set_cell_text(table.rows[0].cells[index], header, bold=True, color=DARK_BLUE)

    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_text(cells[index], value)

    set_table_geometry(table, widths_dxa)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_callout(doc, label, text, caution=False):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.line_spacing = 1.15

    ppr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "FFF4DF" if caution else LIGHT_FILL)
    ppr.append(shading)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "D39E00" if caution else "2E74B5")
    borders.append(left)
    ppr.append(borders)

    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, color=DARK_BLUE, bold=True)
    set_run_font(paragraph.add_run(text), size=10.5)
    return paragraph


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    set_run_font(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    value = OxmlElement("w:t")
    value.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, value, end])


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(
        header_paragraph.add_run("Локальный анонимизатор  |  Инструкция пользователя"),
        size=9,
        color=MUTED,
    )

    footer = section.footer
    add_page_number(footer.paragraphs[0])


def add_title_page(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(82)
    spacer.paragraph_format.space_after = Pt(0)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    set_run_font(kicker.add_run("РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ"), size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    set_run_font(
        title.add_run("Локальный анонимизатор"),
        size=28,
        color=DARK_BLUE,
        bold=True,
    )

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(46)
    set_run_font(
        subtitle.add_run("Безопасная подготовка текста перед передачей во внешние ИИ-сервисы"),
        size=14,
        color=MUTED,
    )

    metadata = [
        "Платформа: macOS 13 и новее",
        "Режим работы: полностью локальный",
        "Актуально для текущей версии MVP",
    ]
    for item in metadata:
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        set_run_font(paragraph.add_run(item), size=11, color=BODY_COLOR)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(86)
    note.paragraph_format.space_after = Pt(0)
    set_run_font(note.add_run("Июль 2026"), size=10, color=MUTED)
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_page(doc)
    configure_styles(doc)
    bullet_num_id = add_numbering_definition(doc, "bullet", "•", FONT)
    decimal_num_id = add_numbering_definition(doc, "decimal", "%1.", FONT)

    core_properties = doc.core_properties
    core_properties.title = "Инструкция пользователя — Локальный анонимизатор"
    core_properties.subject = "Руководство по работе с приложением"
    core_properties.author = "Платформикс"
    core_properties.keywords = "анонимизация, персональные данные, ИИ, macOS"

    add_title_page(doc)

    doc.add_heading("1. Назначение приложения", level=1)
    doc.add_paragraph(
        "«Локальный анонимизатор» помогает удалить или заменить чувствительные данные "
        "перед тем, как пользователь вставит текст во внешний ИИ-сервис."
    )
    add_callout(
        doc,
        "Главное",
        "исходный текст обрабатывается только на компьютере пользователя и не отправляется в интернет.",
    )

    doc.add_heading("Что умеет приложение", level=2)
    for item in [
        "получать обычный текст из системного буфера обмена;",
        "находить чувствительные данные и заменять их псевдонимами;",
        "сохранять абзацы и переносы строк исходного текста;",
        "позволять вручную проверить и исправить результат;",
        "копировать обезличенный текст для последующей работы.",
    ]:
        add_bullet(doc, item, bullet_num_id)

    doc.add_heading("Поддерживаемые категории", level=2)
    add_table(
        doc,
        ["Категория", "Пример исходных данных", "Результат"],
        [
            ("ФИО", "Александр Абрамян", "[PERSON_001]"),
            ("Организации", "ООО «Ромашка»", "[COMPANY_001]"),
            ("Телефоны", "+7 999 123-45-67", "[PHONE_001]"),
            ("Email", "ivan@example.ru", "[EMAIL_001]"),
            ("IPv4", "192.168.1.10", "[IP_ADDRESS_001]"),
            ("Реквизиты", "ИНН, ОГРН, адрес", "[REQUISITE_001]"),
        ],
        [2100, 3660, 3600],
    )

    doc.add_page_break()
    doc.add_heading("2. Быстрый старт", level=1)
    steps = [
        ("Скопируйте текст", "Выделите нужный текст в почте, документе или другой программе и нажмите Command + C."),
        ("Откройте приложение", "Запустите «Локальный анонимизатор»."),
        ("Вставьте текст", "Нажмите кнопку «Вставить из буфера». Анализ начнётся автоматически."),
        ("Проверьте правила", "При необходимости включите категории, добавьте исключения или правила замены слов."),
        ("Повторите обработку", "Если исходный текст или правила изменились, нажмите «Обезличить» либо Command + Enter."),
        ("Проверьте результат", "Просмотрите правую область и при необходимости отредактируйте текст вручную."),
        ("Скопируйте результат", "Нажмите «Скопировать» и вставьте очищенный текст в нужный сервис."),
    ]
    for title, detail in steps:
        add_step(doc, title, detail, decimal_num_id)

    add_callout(
        doc,
        "Важно",
        "перед передачей текста во внешний сервис всегда просматривайте результат: автоматические правила могут пропустить нестандартные данные.",
        caution=True,
    )

    doc.add_heading("3. Основные области окна", level=1)
    definitions = [
        ("Строка состояния", "Показывает готовность приложения, ход анализа, количество найденных фрагментов и ошибки."),
        ("Настройки правил", "Управляют категориями, исключениями и заменой слов."),
        ("Исходный текст", "Содержит текст, который требуется проверить и обезличить."),
        ("Обезличенный текст", "Содержит результат. Его можно редактировать вручную перед копированием."),
    ]
    for label, detail in definitions:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(6)
        set_run_font(paragraph.add_run(f"{label}. "), bold=True, color=DARK_BLUE)
        set_run_font(paragraph.add_run(detail))

    doc.add_page_break()
    doc.add_heading("4. Настройка правил анонимизации", level=1)

    doc.add_heading("4.1. Категории", level=2)
    doc.add_paragraph(
        "Флажки «ФИО», «Организации», «Телефоны», «Email», «IP» и «Реквизиты» определяют, "
        "какие категории приложение ищет автоматически. Снятый флажок отключает "
        "автоматическую замену соответствующей категории."
    )
    doc.add_paragraph(
        "Категория «Реквизиты» заменяет ИНН, КПП, ОГРН, ОГРНИП, ОКПО, БИК, расчётные "
        "и корреспондентские счета, СНИЛС, паспортные данные и адреса. Реквизит должен "
        "быть указан после соответствующей метки. Для адресов используйте формат "
        "«Адрес: …», «Юридический адрес: …», «Фактический адрес: …» или «Почтовый адрес: …»."
    )
    doc.add_paragraph(
        "Категория «ФИО» распознаёт полные формы «Алексей Майборода», "
        "«Майборода Алексей» и сокращения «А.Ковешников», «А. Майборода», "
        "«Ковешников А.», «А.А. Майборода», «Майборода А.А.». Одно ФИО должно "
        "находиться на одной строке."
    )

    doc.add_heading("4.2. Список исключений", level=2)
    doc.add_paragraph(
        "Исключение сохраняет найденный фрагмент без изменения. Введите слово или фразу "
        "и нажмите «Добавить». Для удаления нажмите значок × рядом с термином."
    )
    add_callout(
        doc,
        "Пример",
        "если добавить «Учебный сервер» в исключения, найденный фрагмент с этой фразой останется в исходном виде.",
    )

    doc.add_heading("4.3. Замена слов", level=2)
    doc.add_paragraph(
        "Термины из этого списка всегда заменяются меткой [COMPANY_XXX], даже если "
        "слово встречается без поясняющего контекста. По умолчанию добавлены:"
    )
    for item in ["СИЛА", "Платформикс", "Базовые решения"]:
        add_bullet(doc, item, bullet_num_id)

    doc.add_paragraph(
        "Чтобы добавить слово или фразу, введите значение в поле и нажмите «Добавить». "
        "Звёздочка разрешена только в конце основы: «Иванов*» заменит «Иванов», "
        "«Иванову» и «Ивановым». Чтобы прекратить замену, удалите термин значком ×."
    )

    doc.add_heading("4.4. Приоритет правил", level=2)
    add_callout(
        doc,
        "Порядок",
        "email, телефон, IP и реквизит сначала защищаются как единый фрагмент. "
        "Затем применяется «Замена слов», которая срабатывает даже при выключенной "
        "категории «Организации» и наличии того же слова в исключениях.",
        caution=True,
    )
    add_table(
        doc,
        ["Настройка", "Результат"],
        [
            ("Email, телефон, IP или реквизит", "Заменяется целиком своей меткой"),
            ("Термин в «Замене слов»", "Всегда заменяется на [COMPANY_XXX]"),
            ("Термин только в исключениях", "Остаётся без изменения"),
            ("Категория отключена", "Автоматический поиск категории не выполняется"),
        ],
        [3420, 5940],
    )

    doc.add_heading("5. Горячие клавиши", level=1)
    add_table(
        doc,
        ["Сочетание", "Действие"],
        [
            ("Command + V", "Вставить текст в активную область"),
            ("Command + Enter", "Повторно обработать исходный текст"),
            ("Command + Shift + C", "Скопировать обезличенный результат"),
            ("Command + A", "Выделить весь текст активной области"),
        ],
        [2700, 6660],
    )

    doc.add_heading("6. Статусы и сообщения", level=1)
    add_table(
        doc,
        ["Сообщение", "Что означает / что делать"],
        [
            ("Вставьте текст", "Добавьте исходный текст вручную или из буфера обмена."),
            ("Анализ текста…", "Дождитесь завершения обработки."),
            ("Готово: найдено фрагментов — N", "Результат сформирован и доступен справа."),
            ("Чувствительные данные не найдены", "Проверьте текст и настройки категорий."),
            ("Исходный текст изменён", "Нажмите «Обезличить» для повторного анализа."),
            ("В буфере обмена нет текста", "Скопируйте обычный текст и повторите вставку."),
        ],
        [3660, 5700],
    )

    doc.add_heading("7. Рекомендации по безопасной работе", level=1)
    for item in [
        "проверяйте результат перед отправкой во внешний ИИ-сервис;",
        "не считайте автоматическую обработку единственной мерой защиты;",
        "при работе с новыми типами данных добавляйте необходимые правила замены слов;",
        "не добавляйте в обычные исключения сведения, которые должны быть скрыты;",
        "если правила изменились, всегда выполняйте обработку повторно.",
    ]:
        add_bullet(doc, item, bullet_num_id)

    doc.add_page_break()
    doc.add_heading("8. Конфиденциальность и ограничения", level=1)
    doc.add_heading("Конфиденциальность", level=2)
    for item in [
        "обработка выполняется локально на компьютере пользователя;",
        "приложение не использует платные сервисы и внешние API;",
        "исходный и обезличенный текст не сохраняются в историю;",
        "технические логи не содержат текст пользователя;",
        "локально сохраняются только настройки категорий и списки правил.",
    ]:
        add_bullet(doc, item, bullet_num_id)

    doc.add_heading("Ограничения", level=2)
    doc.add_paragraph(
        "Распознавание выполняется правилами, эвристиками и системными возможностями macOS. "
        "Приложение может пропустить редкое имя, нестандартный формат реквизитов или внутреннее "
        "название проекта. Для повторяющихся внутренних названий используйте блок «Замена слов». "
        "Полная и сокращённая запись одного человека могут получить разные псевдонимы."
    )

    doc.add_heading("Не поддерживается в MVP", level=2)
    for item in [
        "загрузка и сохранение документов;",
        "история обработанных текстов;",
        "отправка результата непосредственно в ИИ-сервис;",
        "гарантированное распознавание всех нестандартных формулировок.",
    ]:
        add_bullet(doc, item, bullet_num_id)

    doc.add_page_break()
    doc.add_heading("9. Частые вопросы", level=1)
    faq = [
        ("Можно ли отредактировать результат?", "Да. Правая текстовая область доступна для ручного редактирования."),
        ("Почему кнопка «Скопировать» неактивна?", "Сначала выполните обработку актуального исходного текста."),
        ("Что происходит после изменения правил?", "Текущий результат становится устаревшим. Запустите обработку повторно."),
        ("Что сильнее: исключение или «Замена слов»?", "Правило из блока «Замена слов» имеет приоритет над исключением. Email, телефон, IP и реквизит при включённой категории защищаются целиком."),
        ("Сохраняется ли введённый текст после закрытия?", "Нет. Сохраняются только настройки правил."),
        ("Почему адрес не был заменён?", "Укажите его после метки «Адрес:», «Юридический адрес:», «Фактический адрес:» или «Почтовый адрес:»."),
        ("Нужен ли интернет?", "Нет. Основной сценарий работает полностью локально."),
    ]
    for question, answer in faq:
        doc.add_heading(question, level=2)
        paragraph = doc.add_paragraph(answer)
        paragraph.paragraph_format.space_after = Pt(6)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
